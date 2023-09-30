import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from docx2pdf import convert # pip install docx2pdf
import copy
import pandas as pd
import docx
from pathlib import Path
from docx.shared import Pt
import locale


class SimSalabim:
    locale.setlocale(locale.LC_MONETARY, '')
    
    def __init__(self, documento:list, base_dados: pd.DataFrame, criarPDF:bool = False, diretorio_destino: Path = f'Documentos Criados - SimSalabim', delimitador:str = "==", fonte_letra:str = 'Times New Roman'):
        self.base_dados = base_dados
        
        self.diretorio_para_salvar = diretorio_destino
        self.nome_documentos:list = documento
        self.doc_base = docx.Document(documento)
        self.fonte_letra = fonte_letra
        
        self.criar_pdf = criarPDF
        
        self.relacao_dataframe = self.dados_dataframe()
        self.relacao_individual_funcionario:list[dict] = list()

        self.delimitador = delimitador
        self.regex_delimitados = list()
        
        self.analisar_delimitadores()
        
        self.relacao_informacoes()
        self.criar_documento()
        self.transformar_pdf()
        
        
    # @staticmethod
    # def adaptacao_salario(valor:float):
    #     return locale.currency(valor, grouping=True) 
    
        
    def transformar_pdf(self):
        if self.criar_pdf:
            diretorio_de_entrada = Path(f"{self.diretorio_para_salvar}/")
            arquivos_docx = [os.path.join(diretorio_de_entrada, arquivo) for arquivo in os.listdir(diretorio_de_entrada) if arquivo.endswith(".docx")]

            for arquivo_docx in arquivos_docx:
                convert(arquivo_docx)

        
    # def dados_dataframe(self):
    #     pl_df = pd.read_excel(self.base_dados)
    #     pl_df = pl_df.sort_values(['NOME'])
    #     pl_df = pl_df.dropna(how='any')
    #     return pl_df
        
    
    def relacao_informacoes(self) -> tuple[dict, str]:
        for index, row in self.relacao_dataframe.iterrows():
            # nome = row['NOME']
            # movimentacao = row['MOVIMENTACAO']
            # funcao = row['Cargo plano salario']
            # salario = row['Tabela Salarial']
            # ctps = row['CTPS']
            
            relacao:dict = {
                '==NOME==':nome,
                '==CTPS==':ctps,
                '==FUNCAO==':funcao,
                '==SALARIO==':self.adaptacao_salario(salario),
            }

            self.relacao_individual_funcionario.append(relacao)
            

    def analisar_delimitadores(self):
        regex = f'{self.delimitador}([^=]+){self.delimitador}'

        for paragrafo in self.doc_base.paragraphs:
            texto = paragrafo.text
            
            correspondencias = re.findall(regex, texto)
            self.regex_delimitados.extend(correspondencias)            
        
                    
    def criar_documento(self, dados:dict, nome_doc:str):
        doc_base_dc = copy.deepcopy(self.doc_base)   

        for paragrafo in doc_base_dc.paragraphs:
            for key, value in dados.items():
                if key in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(key, str(value))
                    for run in paragrafo.runs:
                        run.font.name = self.fonte_letra
                        

        for table in doc_base_dc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in dados.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                            cell.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.name = self.fonte_letra

        doc_base_dc.save(Path(f'{self.diretorio_para_salvar}/{dados.get("==NOME==")} - {nome_doc}.docx'))            
        
        
        
if __name__ =='__main__':
    app = SimSalabim(documento=r'MOVIMENTAÇÃO DE PESSOAL - MUDANÇA DE FUNÇÃO.docx', criarPDF=False, base_dados=r'base.xlsx')
    
## CRIAR ALGORITMO PARA ANALISAR UM DOCX E VER, DEFININDO O DELIMITADOR ANTES "==", QUAIS SAO. DEPOIS, CRIAR UM DICT CONTENDO O DELIMITADOR E RESPECTIVAMENTE OS DADOS QUE POSSUEM O MESMO NOME NA COLUNA DA BASE.XLSX
